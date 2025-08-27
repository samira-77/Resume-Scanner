import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from io import BytesIO

# --- Page Configuration ---
st.set_page_config(page_title="Create Resume", layout="wide")

# --- Custom Styles ---
st.markdown("""
    <style>
    html, body, [data-testid="stApp"] {
        background: linear-gradient(to right, #141e30, #243b55);
        color: white;
        font-family: 'Segoe UI', sans-serif;
    }
    .title {
        font-size: 4rem;
        font-weight: bold;
        color: #00eaff;
        text-shadow: 0 0 15px #00eaff;
        text-align: center;
        margin-bottom: 2rem;
    }
    label {
        font-weight: bold;
        color: #00eaff;
    }
    .stTextInput > div > input, .stTextArea > div > textarea {
        background-color: rgba(0, 0, 0, 0.4);
        color: white;
    }
    .stButton > button {
        background-color: #00eaff;
        color: black;
        font-weight: bold;
        border: none;
        border-radius: 10px;
        padding: 0.7rem 1.5rem;
        transition: all 0.3s ease;
        margin-top: 1.5rem;
    }
    .stButton > button:hover {
        background-color: #0288d1;
        color: white;
        box-shadow: 0 0 10px #00eaff;
    }
    </style>
""", unsafe_allow_html=True)

# --- Title ---
st.markdown('<div class="title">Create Your Resume</div>', unsafe_allow_html=True)

# --- Resume Form ---
with st.form("resume_form"):
    col1, col2 = st.columns(2)

    with col1:
        name = st.text_input("Full Name")
        email = st.text_input("Email Address")
        phone = st.text_input("Phone Number")
        linkedin = st.text_input("LinkedIn URL")
        portfolio = st.text_input("Portfolio URL (Optional)")

    with col2:
        summary = st.text_area("Professional Summary", height=150)
        skills = st.text_area("Key Skills (comma-separated)", height=100)
        experience = st.text_area("Work Experience (job title, company, achievements)", height=200)
        education = st.text_area("Education (degree, college, year)", height=100)
        projects = st.text_area("Projects (if any)", height=100)

    # Template Selection
    template_choice = st.selectbox(
        "Choose Resume Template Style:",
        ("Modern", "Elegant", "Compact")
    )

    submitted = st.form_submit_button("Generate Resume")

# --- Helper Function ---
def create_resume_docx(name, email, phone, linkedin, portfolio, summary, skills, experience, education, projects, template_choice):
    doc = Document()

    # --- Set Page Margins ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- Header ---
    header = doc.add_paragraph()
    run = header.add_run(name)
    run.font.size = Pt(22)
    run.bold = True
    if template_choice == "Modern":
        run.font.color.rgb = RGBColor(0, 153, 255)
    elif template_choice == "Elegant":
        run.font.color.rgb = RGBColor(0, 0, 0)
    else:
        run.font.color.rgb = RGBColor(80, 80, 80)
    header.alignment = 1

    # --- Contact Info ---
    contact_info = f"{email} | {phone}"
    if linkedin:
        contact_info += f" | LinkedIn: {linkedin}"
    if portfolio:
        contact_info += f" | Portfolio: {portfolio}"
    p = doc.add_paragraph(contact_info)
    p.alignment = 1

    # --- Section Helper ---
    def add_section(title, content, bullet=False):
        doc.add_paragraph()
        heading = doc.add_paragraph(title)
        heading.runs[0].bold = True
        heading.runs[0].font.size = Pt(14)
        heading.runs[0].font.color.rgb = RGBColor(0, 102, 204) if template_choice == "Modern" else RGBColor(0, 0, 0)

        if bullet:
            for item in content.split(','):
                doc.add_paragraph(item.strip(), style='List Bullet')
        else:
            doc.add_paragraph(content)

    # --- Resume Sections ---
    add_section("Professional Summary", summary)
    add_section("Key Skills", skills, bullet=True)
    add_section("Work Experience", experience)
    add_section("Education", education)
    if projects.strip():
        add_section("Projects", projects)

    # --- Save ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- Generate Resume and Download ---
if submitted:
    docx_file = create_resume_docx(
        name, email, phone, linkedin, portfolio,
        summary, skills, experience, education,
        projects, template_choice
    )
    st.success("🎉 Resume Created Successfully!")

    st.download_button(
        label="📄 Download Resume as Word (.docx)",
        data=docx_file,
        file_name=f"{name.replace(' ', '_')}_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# --- Back to Home Button ---
st.markdown("""
    <a href="/" style="text-decoration: none;">
        <button style="
            background-color: #00eaff;
            color: black;
            font-weight: bold;
            border: none;
            border-radius: 10px;
            padding: 0.7rem 1.5rem;
            transition: all 0.3s ease;
            margin-top: 1.5rem;
            cursor: pointer;
        ">⬅ Back to Home</button>
    </a>
""", unsafe_allow_html=True)
